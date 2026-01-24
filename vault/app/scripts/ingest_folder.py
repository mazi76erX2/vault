import argparse
from pathlib import Path

from sqlalchemy import func
from sqlalchemy.orm import Session

from app.core.config import settings
from app.core.database import SessionLocal  # Standardize on core database
from app.integrations.ollama_client import embed
from app.models.kb import KBChunk, KBDocument
from app.services.rag_service import chunk_text


def read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def read_pdf(path: Path) -> str:
    from PyPDF2 import PdfReader

    reader = PdfReader(str(path))
    parts = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return "\n".join(parts)


def read_docx(path: Path) -> str:
    import docx

    d = docx.Document(str(path))
    return "\n".join(p.text for p in d.paragraphs)


def extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".txt":
        return read_txt(path)
    if ext == ".pdf":
        return read_pdf(path)
    if ext == ".docx":
        return read_docx(path)
    return ""


def upsert_doc(db: Session, title: str, sourcefile: str, chunks: list[str], full_text: str) -> None:
    # 1. Create or get the document record
    kb_doc = db.query(KBDocument).filter(KBDocument.sourcefile == sourcefile).first()
    if not kb_doc:
        kb_doc = KBDocument(
            title=title,
            sourcefile=sourcefile,
            content=full_text,
            accesslevel=1,
            # Populate parent TSVector for Hybrid Search
            tsv=func.to_tsvector('english', full_text)
        )
        db.add(kb_doc)
        db.flush() # Get ID
    else:
        # Update existing
        kb_doc.content = full_text
        kb_doc.tsv = func.to_tsvector('english', full_text)

    # 2. Clear old chunks
    db.query(KBChunk).filter(KBChunk.doc_id == kb_doc.id).delete()

    # 3. Create new chunks with embeddings and TSVector
    rows: list[KBChunk] = []
    for i, ch in enumerate(chunks, start=1):
        embedding = embed(ch)
        rows.append(
            KBChunk(
                doc_id=kb_doc.id,
                chunk_index=i,
                title=title,
                sourcefile=sourcefile,
                content=ch,
                accesslevel=1,
                embedding=embedding,
                # Explicitly populate TSVector for Hybrid Search per chunk
                tsv=func.to_tsvector('english', ch)
            )
        )

    db.add_all(rows)
    db.commit()


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--folder", required=True, help="Folder with .txt/.pdf/.docx files")
    parser.add_argument("--chunk-size", type=int, default=settings.KB_CHUNK_SIZE)
    parser.add_argument("--chunk-overlap", type=int, default=settings.KB_CHUNK_OVERLAP)
    args = parser.parse_args()

    folder = Path(args.folder)
    if not folder.exists():
        raise SystemExit(f"Folder not found: {folder}")

    db = SessionLocal()
    try:
        for path in folder.rglob("*"):
            if not path.is_file():
                continue
            if path.suffix.lower() not in {".txt", ".pdf", ".docx"}:
                continue

            text = extract_text(path).strip()
            if not text:
                continue

            # Uses Semantic Chunking by default via rag_service (strategy='semantic')
            chunks = chunk_text(text, chunk_size=args.chunk_size, overlap=args.chunk_overlap, strategy="semantic")
            if not chunks:
                continue

            title = path.name
            sourcefile = str(path)

            upsert_doc(db, title=title, sourcefile=sourcefile, chunks=chunks, full_text=text)
            print(f"Ingested {path} ({len(chunks)} chunks)")

    except Exception as e:
        print(f"Error during ingestion: {e}")
    finally:
        db.close()


if __name__ == "__main__":
    main()
